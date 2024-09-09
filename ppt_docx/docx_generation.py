import os
import hashlib
import time
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict
from env import get_app_root
from project.file_manage import clear_files_by_timediff
from project.schedule import get_scheduler

_OUTPUT_DIR_DOCX = os.path.join(get_app_root(), "data/cache/docx")

# 如果文件夹路径不存在，先创建
if not os.path.exists(_OUTPUT_DIR_DOCX):
    os.makedirs(_OUTPUT_DIR_DOCX)

def get_file_path_docx(text):
    """生成唯一的文件路径"""
    file_name = hashlib.sha256(text.encode("utf-8")).hexdigest()  # 可以使用uuid替代
    return os.path.join(_OUTPUT_DIR_DOCX, f"{file_name}.docx")

def generate_docx_content(docx_content: Dict) -> str:
    """生成 docx 文件"""
    document = Document()

    # Word 标题
    document.add_heading(docx_content['title'], 0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 页内容
    print(f'总共 {len(docx_content["sections"])} 个章节')
    for i, section in enumerate(docx_content['sections']):
        print(f'生成第 {i + 1} 章节: {section["heading"]}')
        document.add_heading(section['heading'], level=1)
        for paragraph in section['paragraphs']:
            document.add_heading(paragraph['heading'], level=2)
            p = document.add_paragraph(paragraph['content'])
            p.style.font.size = Pt(12)

    _output_file = get_file_path_docx(str(time.time()))
    document.save(_output_file)

    return _output_file

def clear_docx_cache():
    """清理生成的 docx 缓存"""
    try:
        clear_files_by_timediff(_OUTPUT_DIR_DOCX, datetime.timedelta(minutes=10).seconds)
    except Exception as e:
        pass

# 定时清理缓存的 job
get_scheduler().add_job(clear_docx_cache, "interval", seconds=10, id="clear_docx_cache")
